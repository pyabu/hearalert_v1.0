import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def configure_professional_styles(doc):
    # Set Margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styling Normal text
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Justified alignment
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.space_after = Pt(12)
    style.paragraph_format.line_spacing = 1.15

    # Styling Headings
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Arial'
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 51, 102) # Dark blue
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Arial'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 51, 102)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)
    
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Arial'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(50, 50, 50)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)

def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    configure_professional_styles(doc)

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    
    in_table = False
    table_lines = []

    for line in lines:
        if line.startswith('---'):
            continue
            
        # Handle Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            # Center the title if it is the very first heading
            if "Project Report:" in line:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
        
        # Handle lists
        elif line.startswith('*   ') or line.startswith('- '):
            p = doc.add_paragraph(line[4:], style='List Bullet')
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(text, style='List Number')
            
        # Handle Tables
        elif line.startswith('|'):
            if '---' in line:
                continue 
            if not in_table:
                in_table = True
                table_lines = [line]
            else:
                table_lines.append(line)
        else:
            if in_table:
                if len(table_lines) > 0:
                    cols = len([c for c in table_lines[0].split('|') if c.strip()])
                    table = doc.add_table(rows=len(table_lines), cols=cols)
                    table.style = 'Light Shading Accent 1' # Professional built-in
                    
                    for row_idx, t_line in enumerate(table_lines):
                        cells = [c.strip() for c in t_line.split('|') if c.strip()]
                        for col_idx, cell_text in enumerate(cells):
                            if col_idx < cols:
                                cell = table.cell(row_idx, col_idx)
                                cell.text = cell_text.replace('**', '')
                                
                                # Make header bold
                                if row_idx == 0:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.bold = True
                in_table = False
                table_lines = []
                
            # Check for Images
            img_match = re.match(r'!\[.*?\]\((.*?)\)', line.strip())
            if img_match:
                img_path = img_match.group(1)
                full_img_path = os.path.join(os.path.dirname(md_path), img_path)
                if os.path.exists(full_img_path):
                    # Center align the image
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run()
                    r.add_picture(full_img_path, width=Inches(6.0))
                else:
                    doc.add_paragraph(f"[Missing Image: {img_path}]")
                continue
                
            # Normal paragraph
            if line.strip() != '':
                p = doc.add_paragraph()
                
                # Check for center styling specifically for subtitle
                if "**AI-Powered Environmental Awareness" in line:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    line = line.replace('**', '') # strip asterisks for subtitle
                    r = p.add_run(line)
                    r.bold = True
                    r.font.size = Pt(13)
                    r.font.color.rgb = RGBColor(100, 100, 100)
                    continue

                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)

    doc.save(docx_path)
    print(f"Successfully converted to {docx_path}")

md_file = "/Users/abusaleem/hearalert-v1.1/hearalertt/Second_Review_Report.md"
docx_file = "/Users/abusaleem/hearalert-v1.1/hearalertt/Second_Review_Report.docx"

if os.path.exists(md_file):
    convert_md_to_docx(md_file, docx_file)
else:
    print("Markdown file not found.")
