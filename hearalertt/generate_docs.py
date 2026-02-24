from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.name = 'Arial'

def add_paragraph(doc, text, bold=False, indent=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    if indent > 0:
        p.paragraph_format.left_indent = Inches(indent)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_bullet(doc, text, indent=0.5):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(indent)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Create a new Document
doc = Document()

# Add Title
title = doc.add_heading('HearAlert Audio Classification Pipeline', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = 'Arial'

doc.add_paragraph() # Spacing

# Overview
add_heading(doc, 'Overview', level=1)
add_paragraph(doc, "The HearAlert audio training pipeline is designed to build a robust, real-time audio classification model capable of running efficiently on edge devices (like mobile phones). It leverages a Transfer Learning pipeline based on YAMNet coupled with a custom deep neural network (DNN) head and rigorous on-the-fly data augmentation.\n\nThe complete pipeline is handled by train_audio_model.py and download_realtime_datasets.py.")

doc.add_paragraph() # Spacing

# 1. Data Aggregation & Splitting
add_heading(doc, '1. Data Aggregation & Splitting', level=2)
add_paragraph(doc, "The pipeline aggregates over 19,000 files to form a diverse and robust training dataset covering 33 distinct environmental, emergency, and human sound categories.")

add_paragraph(doc, "Data Sources:", bold=True, indent=0.25)
add_bullet(doc, "Base ESC-50 dataset")
add_bullet(doc, "Previously processed and augmented raw data")
add_bullet(doc, "Synthesized real-time sounds (e.g., Smoke Alarms, Car Alarms, Microwaves, Water Running) generated to fill gaps where real-world data was sparse.")

add_paragraph(doc, "Splitting:", bold=True, indent=0.25)
add_bullet(doc, "80% Training Set: Used to update model weights.")
add_bullet(doc, "10% Validation Set: Used to tune hyperparameters and trigger early stopping.")
add_bullet(doc, "10% Testing Set: Held out to verify final unseen accuracy.")

# 2. Feature Extraction (Transfer Learning)
add_heading(doc, '2. Feature Extraction (Transfer Learning)', level=2)
add_paragraph(doc, "Instead of feeding raw audio directly into a massive ConvNet, we use Google’s pre-trained YAMNet (which is already highly trained on the AudioSet corpus).")
add_bullet(doc, "Audio Processing: Incoming audio waveforms are padded/trimmed to precisely 1 second at a 16kHz sample rate.", indent=0.5)
add_bullet(doc, "YAMNet Embeddings: The waveform is passed through YAMNet, which extracts a condensed, highly informative 1024-dimensional feature embedding. This approach drastically reduces the computational load and time required to train the custom head.", indent=0.5)

# 3. On-The-Fly Data Augmentation
add_heading(doc, '3. On-The-Fly Data Augmentation', level=2)
add_paragraph(doc, "To simulate real-world conditions (different rooms, background noises, varied hardware microphones) and prevent overfitting, the training samples undergo dynamic augmentation before feature extraction. During training, two extra augmented versions are generated per original sample using:")
add_bullet(doc, "Background Noise Injection (50% chance): Adds Gaussian noise to simulate poor microphone quality or environmental noise.", indent=0.5)
add_bullet(doc, "Time Shifting (50% chance): Rolls the audio array to change when the sound event occurs in the 1-second window.", indent=0.5)
add_bullet(doc, "Volume/Gain Variation (50% chance): Multiplies the signal by a random gain between 0.7x and 1.3x.", indent=0.5)
add_bullet(doc, "Pitch & Speed Modification (30% chance): Slightly speeds up or slows down the sound event to alter its pitch perception.", indent=0.5)

# 4. Custom Model Architecture
add_heading(doc, '4. Custom Model Architecture', level=2)
add_paragraph(doc, "A custom Deep Neural Network (DNN) is trained on top of the YAMNet embeddings to classify the 33 specific HearAlert categories.\n\nThe architecture incorporates Batch Normalization and Dropout to ensure stable, generalized learning:")
add_bullet(doc, "Input Layer: 1024 parameters (YAMNet embeddings)")
add_bullet(doc, "Dense Block 1: 512 units → Batch Normalization → ReLU Activation → Dropout (40%)")
add_bullet(doc, "Dense Block 2: 256 units → Batch Normalization → ReLU Activation → Dropout (30%)")
add_bullet(doc, "Dense Block 3: 128 units → Batch Normalization → ReLU Activation → Dropout (20%)")
add_bullet(doc, "Output Layer: 33 units with Softmax activation for multi-class probability prediction.")

# 5. Advanced Training Strategies
add_heading(doc, '5. Advanced Training Strategies', level=2)
add_paragraph(doc, "The training regimen employs several advanced machine learning techniques to squeeze maximal accuracy out of the dataset:")
add_bullet(doc, "Class Balancing: Uses compute_class_weight to automatically apply higher priority/weight to minority classes, preventing the model from becoming biased toward classes with overwhelmingly large sample sizes.")
add_bullet(doc, "Cosine Decay with Warmup: The Adam optimizer utilizes a learning rate schedule that slowly warms up for the first 5 epochs, and then follows a cosine decay curve to smoothly settle into the optimal minima.")
add_bullet(doc, "Label Smoothing: Uses SparseCategoricalCrossentropy combined with a smooth loss landscape configuration to prevent the model from becoming overconfident in noisy real-world data.")

add_paragraph(doc, "Dynamic Callbacks:", bold=True, indent=0.25)
add_bullet(doc, "Early Stopping: Halts training if validation accuracy ceases to improve for 15 consecutive epochs.", indent=0.75)
add_bullet(doc, "ReduceLROnPlateau: Dynamically reduces the learning rate by half if the validation loss plateaus.", indent=0.75)
add_bullet(doc, "Model Checkpoint: Saves only the highest-performing iteration (highest val_accuracy) during the 100-epoch run.", indent=0.75)

# 6. Edge Optimization (INT8 Quantization)
add_heading(doc, '6. Edge Optimization (INT8 Quantization)', level=2)
add_paragraph(doc, "To prepare the model for rapid inference on Android/iOS without draining the battery:")
add_bullet(doc, "The finest iteration of the model is converted using TensorFlow Lite (TFLiteConverter).")
add_bullet(doc, "Full-Integer Quantization: A representative dataset (drawn from the validation set) is fed into the converter to calibrate the dynamic range. The 32-bit floating-point weights and activations are fully quantized down to 8-bit integers (INT8).")
add_bullet(doc, "Result: The model shrinks from several megabytes down to ~704 KB, achieving a footprint small enough to ship directly in the mobile app bundle while keeping inference latency in the low milliseconds.")

# Outputs
add_heading(doc, 'Outputs', level=2)
add_paragraph(doc, "At the successful conclusion of the pipeline, the following files are updated in mobile_app/assets/models/:")
add_bullet(doc, "1. hearalert_classifier.tflite – The final, edge-ready quantized model.")
add_bullet(doc, "2. hearalert_labels.txt – The ordered list of the 33 categories matching the neural network outputs.")
add_bullet(doc, "3. categories_config.json – A metadata map combining categories with UI configurations (display names, priority levels, vibration patterns, etc.) for app integration.")


# Save the document
output_path = '/Users/abusaleem/Hearalert-version-1.1/hearalert-1.0.1/hearalertt/AUDIO_TRAINING_PIPELINE.docx'
doc.save(output_path)
print(f"Document successfully created at {output_path}")
