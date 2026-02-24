#!/usr/bin/env python3
"""
Generate HearAlert Metrics Documentation as a Word (.docx) file.
Usage: python3 generate_metrics_doc.py
"""

import subprocess, sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx", "-q"], check=True)
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

from pathlib import Path

OUT_PATH = Path(__file__).parent / "HearAlert_Metrics_Documentation.docx"

# ─── Helpers ──────────────────────────────────────────────────────────────────
def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p


def body(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10.5)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Header background color (dark indigo)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "4F46E5")
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        fill = "F5F3FF" if r_idx % 2 == 0 else "FFFFFF"
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)

    return table


# ─── Build Document ───────────────────────────────────────────────────────────
def build_doc():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Title ─────────────────────────────────────────────────────────────────
    title = doc.add_heading("HearAlert Audio Classifier", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(79, 70, 229)  # Indigo

    sub = doc.add_paragraph("Training Metrics — Purpose & Results")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(13)
    sub.runs[0].font.color.rgb = RGBColor(107, 91, 149)

    doc.add_paragraph("Generated: February 23, 2026").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── Section 1: Overview ───────────────────────────────────────────────────
    heading(doc, "1. Overview", 1, color=(79, 70, 229))
    body(doc, (
        "HearAlert is an assistive app for deaf and hard-of-hearing users that detects "
        "environmental sounds in real-time and alerts users through vibration patterns. "
        "The audio classifier was trained on 31,386 WAV files across 33 sound categories "
        "using a YAMNet-based transfer learning pipeline."
    ))
    doc.add_paragraph()

    # ── Section 2: Training Results ───────────────────────────────────────────
    heading(doc, "2. Training Results", 1, color=(79, 70, 229))
    add_table(doc,
        headers=["Metric", "Value"],
        rows=[
            ["Training Samples",      "25,109 files"],
            ["Validation Samples",    "3,139 files"],
            ["Test Samples",          "3,138 files"],
            ["Best Train Accuracy",   "62.54%"],
            ["Best Validation Accuracy", "62.42%"],
            ["Test Accuracy (honest)", "62.00%"],
            ["Model Size",            "705 KB (TFLite quantized)"],
            ["Categories",            "33 sound classes"],
        ]
    )
    doc.add_paragraph()

    # ── Section 3: Metrics Explained ─────────────────────────────────────────
    heading(doc, "3. Metrics Used — Why Each Was Chosen", 1, color=(79, 70, 229))

    metrics = [
        (
            "3.1  Accuracy",
            "Accuracy measures what percentage of all sounds are classified correctly overall.",
            "HearAlert needs a single number to compare model versions quickly after each "
            "training run. At 62%, 6 out of every 10 sounds are identified correctly. "
            "Since there are 33 categories, random guessing would only achieve ~3% — "
            "so 62% represents strong performance.",
            [
                ["Training",   "62.54%"],
                ["Validation", "62.42%"],
                ["Test",       "62.00%"],
            ]
        ),
        (
            "3.2  Precision",
            "Precision measures: of all times the model said 'this sound is X', how often was it actually X?",
            "A deaf user receives vibration alerts. If the phone vibrates saying 'Siren!' when "
            "it is actually traffic noise, that is a false alarm. Repeated false alarms erode user "
            "trust — users will start ignoring all alerts. High precision means alerts are reliable. "
            "For example, Siren achieves 99.9% precision — virtually every siren alert is real.",
            [
                ["Siren",              "99.9% ✅"],
                ["Smoke Alarm",        "100.0% ✅"],
                ["Fire Alarm",         "96.8% ✅"],
                ["Car Horn",           "36.7% ⚠️ (needs improvement)"],
                ["Cat Meow",           "9.0% ❌ (causes false alerts)"],
            ]
        ),
        (
            "3.3  Recall",
            "Recall measures: of all real occurrences of sound X, how many did the model detect?",
            "A deaf person cannot hear missed sounds. If the app misses a fire alarm or baby "
            "crying, the user is in danger. Missing a critical detection is worse than a false alarm. "
            "High recall ensures the app catches real-world sounds reliably.",
            [
                ["Water Running",   "100.0% ✅"],
                ["Siren",          "98.0% ✅"],
                ["Footsteps",      "94.7% ✅"],
                ["Door Knock",     "17.3% ❌ (misses most knocks)"],
                ["Glass Breaking", "23.0% ❌ (misses most breaks)"],
            ]
        ),
        (
            "3.4  F1-Score",
            "F1-Score is the harmonic mean of Precision and Recall, combining both into one balanced number.",
            "Some categories have high precision but low recall (or vice versa). F1-Score captures "
            "both in a single value. It is the most useful metric for identifying which sound categories "
            "need more training data. For HearAlert, a category with low F1 either creates too many "
            "false alarms (low precision) or misses too many real sounds (low recall).",
            [
                ["Siren",            "98.9% ✅"],
                ["Background",       "94.5% ✅"],
                ["Footsteps",        "94.7% ✅"],
                ["Smoke Alarm",      "100.0% ✅"],
                ["Glass Breaking",   "23.2% ❌ — needs more data"],
                ["Cat Meow",         "16.0% ❌ — needs more data"],
            ]
        ),
        (
            "3.5  Class Weights (Balanced)",
            "Class weights penalise the model more heavily for misclassifying under-represented categories.",
            "The dataset has imbalanced categories: background has 1,557 files while alarm_clock has "
            "only 540 files. Without class weights, the model would learn to favour predicting "
            "'background' for everything — getting easy accuracy by ignoring small classes. "
            "Class weights ensure every sound category receives equal training attention, "
            "regardless of file count.",
            [
                ["Background (largest)",  "1,557 files — lower weight"],
                ["Alarm Clock (smallest)", "540 files — higher weight"],
                ["Effect", "Model treats all 33 categories fairly"],
            ]
        ),
        (
            "3.6  Validation Loss (Early Stopping)",
            "Validation loss measures how well the model performs on unseen data. Early stopping halts "
            "training when validation loss stops improving.",
            "Without early stopping, the model would memorise training files and fail on real-world audio "
            "(overfitting). The model stopped at epoch 67 out of 80 — before overfitting began. "
            "This ensures the deployed TFLite model generalises to new sounds the user encounters "
            "in the real world, not just sounds from the training dataset.",
            [
                ["Patience",       "12 epochs"],
                ["Best Epoch",     "67"],
                ["Final Val Loss", "1.207 (improving)"],
                ["Result", "Model weights restored to best checkpoint"],
            ]
        ),
    ]

    for metric_title, definition, reason, examples in metrics:
        heading(doc, metric_title, 2, color=(109, 40, 217))

        p = doc.add_paragraph()
        run = p.add_run("Definition:  ")
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(definition).font.size = Pt(11)

        p2 = doc.add_paragraph()
        run2 = p2.add_run("Why Used in HearAlert:  ")
        run2.bold = True
        run2.font.size = Pt(11)
        p2.add_run(reason).font.size = Pt(11)

        doc.add_paragraph("Examples:", style="Normal").runs[0].bold = True
        add_table(doc, headers=["Category / Parameter", "Value / Result"], rows=examples)
        doc.add_paragraph()

    # ── Section 4: Per-Category Results ───────────────────────────────────────
    heading(doc, "4. Per-Category Accuracy (Test Set)", 1, color=(79, 70, 229))
    body(doc, "Results on the held-out test set (file-level split, no data leakage):")
    doc.add_paragraph()

    cat_rows = [
        ["🟢 siren",            "100%", "99.9%", "98.0%", "98.9%"],
        ["🟢 smoke_alarm",      "100%", "100%",  "100%",  "100%"],
        ["🟢 alarm_clock",      "100%", "100%",  "100%",  "100%"],
        ["🟢 car_alarm",        "100%", "100%",  "100%",  "100%"],
        ["🟢 knock_knock",      "100%", "31.1%", "100%",  "47.4%"],
        ["🟢 microwave_beep",   "100%", "99.6%", "100%",  "99.8%"],
        ["🟢 water_running",    "100%", "99.2%", "100%",  "99.6%"],
        ["🟢 footsteps",        "97.8%","94.7%", "94.7%", "94.7%"],
        ["🟢 background",       "98.1%","96.7%", "92.5%", "94.5%"],
        ["🟢 speech",           "96.2%","97.0%", "77.4%", "86.1%"],
        ["🟢 door_creaking",    "90.2%","76.7%", "87.2%", "81.6%"],
        ["🟡 fire_alarm",       "69.3%","96.8%", "69.3%", "80.8%"],
        ["🟡 doorbell",         "61.3%","99.4%", "63.8%", "77.7%"],
        ["🟡 chainsaw",         "60.2%","97.4%", "55.4%", "70.6%"],
        ["🟡 dog_bark",         "57.0%","38.8%", "37.0%", "37.9%"],
        ["🟡 keyboard_typing",  "57.9%","88.4%", "44.0%", "58.7%"],
        ["🟡 breathing",        "54.7%","89.4%", "46.8%", "61.5%"],
        ["🟡 vacuum_cleaner",   "53.8%","22.1%", "52.0%", "31.0%"],
        ["🟡 phone_ring",       "53.8%","61.9%", "37.9%", "47.0%"],
        ["🔴 cat_meow",         "45.7%","9.0%",  "68.1%", "16.0%"],
        ["🔴 coughing",         "45.2%","20.2%", "29.1%", "23.9%"],
        ["🔴 clock_tick",       "45.7%","89.5%", "34.4%", "49.7%"],
        ["🔴 baby_cry",         "43.3%","90.5%", "42.6%", "57.9%"],
        ["🔴 airplane",         "42.6%","83.6%", "38.9%", "53.1%"],
        ["🔴 train",            "41.5%","87.2%", "40.6%", "55.4%"],
        ["🔴 helicopter",       "51.6%","93.3%", "38.5%", "54.5%"],
        ["🔴 thunderstorm",     "49.5%","91.7%", "39.5%", "55.2%"],
        ["🔴 car_horn",         "40.4%","36.7%", "36.0%", "36.4%"],
        ["🔴 washing_machine",  "39.4%","89.1%", "38.5%", "53.8%"],
        ["🔴 gunshot_firework", "40.0%","48.9%", "36.6%", "41.9%"],
        ["🔴 door_knock",       "25.2%","63.8%", "17.3%", "27.2%"],
        ["🔴 traffic",          "32.8%","96.8%", "25.8%", "40.8%"],
        ["🔴 glass_breaking",   "38.0%","23.4%", "23.0%", "23.2%"],
    ]
    add_table(doc,
        headers=["Category", "Test Acc", "Precision", "Recall", "F1-Score"],
        rows=cat_rows
    )
    doc.add_paragraph()

    # ── Section 5: Summary ────────────────────────────────────────────────────
    heading(doc, "5. Summary", 1, color=(79, 70, 229))
    add_table(doc,
        headers=["Metric", "The Real Question It Answers"],
        rows=[
            ["Accuracy",      "Is the model getting better between training runs?"],
            ["Precision",     "Will users trust the vibration alerts?"],
            ["Recall",        "Will users be kept safe from missed sounds?"],
            ["F1-Score",      "Which categories need more training data?"],
            ["Class Weights", "Is every sound category treated fairly during training?"],
            ["Val Loss",      "Will the model work on real-world audio (not just training files)?"],
        ]
    )

    doc.add_paragraph()
    body(doc, (
        "The categories with lowest F1 scores (cat_meow, glass_breaking, door_knock, traffic) "
        "are acoustically similar to other categories and would benefit from more unique, "
        "high-quality recordings to improve detection accuracy."
    ), italic=True)

    doc.save(str(OUT_PATH))
    print(f"\n✅ Document saved: {OUT_PATH}\n")


if __name__ == "__main__":
    build_doc()
