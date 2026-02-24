from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.name = 'Arial'

def add_paragraph(doc, text, bold=False, italic=False, indent=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if indent > 0:
        p.paragraph_format.left_indent = Inches(indent)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_bullet(doc, text, indent=0.5):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(indent)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Create a new Document
doc = Document()

# Add Title
title = doc.add_heading('HearAlert Project Overview & Architecture', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = 'Arial'

doc.add_paragraph() # Spacing

# 1. Introduction
add_heading(doc, '1. Introduction & Project Scope', level=1)
add_paragraph(doc, "HearAlert is an intelligent, real-time audio classification mobile application aimed at providing critical acoustic awareness. The primary audience includes individuals who are deaf or hard of hearing, as well as users needing spatial awareness in noise-canceling environments. By utilizing an edge-based machine learning model, HearAlert identifies essential environmental, emergency, and human sounds, instantly translating them into visual cues and customizable haptic feedback (vibrations).")

# 2. High-Level System Architecture
add_heading(doc, '2. High-Level System Architecture', level=1)
add_paragraph(doc, "The project is divided into two primary technical pillars:")

add_paragraph(doc, "A. Machine Learning & Audio Pipeline (Python)", bold=True, indent=0.25)
add_paragraph(doc, "A robust backend environment responsible for data aggregation, synthesis, and deep learning. It produces the highly optimized TensorFlow Lite (TFLite) models consumed by the app.", indent=0.5)

add_paragraph(doc, "B. Frontend Mobile Application (Flutter & Dart)", bold=True, indent=0.25)
add_paragraph(doc, "A cross-platform mobile application that handles real-time microphone telemetry, local inference of the TFLite model, user interface (UI) rendering, and system-level alerts (notifications/vibrations).", indent=0.5)

# 3. Part 1: Machine Learning Pipeline
add_heading(doc, '3. Machine Learning Pipeline (Backend)', level=1)
add_paragraph(doc, "The Python-based ML pipeline ensures the edge model is accurate, lightweight, and capable of real-time multi-class tracking.")

add_paragraph(doc, "Key Components:", bold=True)
add_bullet(doc, "Data Aggregation: Merges the base ESC-50 dataset with dynamically synthesized sounds (using oscillators, noise injection, and envelopes) across 33 distinct categories.")
add_bullet(doc, "Data Augmentation: Dynamically applies Pitch Shifting, Time Shifting, and Background Noise Injection during training to ensure robustness against different microphone qualities.")
add_bullet(doc, "Feature Extraction (YAMNet): Utilizes Google's pre-trained YAMNet model to extract 1024-dimensional embeddings from 1-second audio windows, accelerating training and boosting accuracy.")
add_bullet(doc, "Custom Neural Network Head: Classifies the extracted embeddings into categories using Dense layers equipped with BatchNormalization, dropout, and Label Smoothing.")
add_bullet(doc, "Edge Optimization: Compresses the best-performing model utilizing INT8 Full-Integer Quantization, resulting in a ~700KB TFLite model ready for deployment without heavy battery drain.")

# 4. Part 2: Mobile Application Structure
add_heading(doc, '4. Mobile Application Structure (Frontend)', level=1)
add_paragraph(doc, "The Flutter application is built applying a modular structure for maintainability and scalability.")

add_paragraph(doc, "4.1 Key Services (/lib/services/)", bold=True, indent=0.25)
add_bullet(doc, "HearAlertClassifierService: Interfaces directly with the TFLite runtime. Buffers microphone streams, enforces sample rates (16kHz), and triggers asynchronous model inferences.", indent=0.5)
add_bullet(doc, "NotificationService: Manages local system notifications. Keeps the user informed of background alerts when the app is minimized.", indent=0.5)
add_bullet(doc, "AlertService: Maps the incoming classification results to specific, priority-driven haptic vibration patterns configured dynamically via JSON.", indent=0.5)

add_paragraph(doc, "4.2 User Interface (/lib/screens/ & /lib/widgets/)", bold=True, indent=0.25)
add_bullet(doc, "Onboarding Screen: Secures necessary OS-level permissions (Microphone, Notifications) while explaining the app’s functionality.", indent=0.5)
add_bullet(doc, "Home Screen: Displays a real-time listening status, the currently detected sound, and a historical log of recent alerts.", indent=0.5)

add_paragraph(doc, "4.3 State Management (/lib/providers/)", bold=True, indent=0.25)
add_bullet(doc, "SoundProvider: Utilizes the Provider pattern to decouple business logic from the UI. It continuously broadcasts listening states and detection events across the widget tree.", indent=0.5)

# 5. Configuration & Integration
add_heading(doc, '5. Integration & Configuration', level=1)
add_paragraph(doc, "The bridge between the ML Pipeline and the Mobile App is entirely data-driven.")
add_bullet(doc, "categories_config.json: Automatically compiled during the Python training phase, this JSON file carries priorities, display names, and distinct vibration arrays (e.g., [0, 500, 200, 500] for sirens) over to the Dart application.")
add_bullet(doc, "hearalert_classifier.tflite: The model graph itself seamlessly drops into the Flutter /assets directory.")

# Save the document
output_path = '/Users/abusaleem/Hearalert-version-1.1/hearalert-1.0.1/hearalertt/HEARALERT_PROJECT_OVERVIEW.docx'
doc.save(output_path)
print(f"Document successfully created at {output_path}")
