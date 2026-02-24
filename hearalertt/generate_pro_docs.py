from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Initialize Document
doc = Document()

# --- Set Page Margins for Professional Documentation ---
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.name = 'Arial'

def add_paragraph(doc, text, bold=False, italic=False, indent=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if indent > 0:
        p.paragraph_format.left_indent = Inches(indent)
    
    # Professional spacing
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_step(doc, step_num, title, description):
    p = doc.add_paragraph()
    # Step Number
    run_num = p.add_run(f"Step {step_num}: ")
    run_num.font.name = 'Arial'
    run_num.font.size = Pt(11)
    run_num.bold = True
    
    # Step Title
    run_title = p.add_run(title)
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(11)
    run_title.bold = True
    
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    desc_p = doc.add_paragraph()
    run_desc = desc_p.add_run(description)
    run_desc.font.name = 'Arial'
    run_desc.font.size = Pt(11)
    
    desc_p.paragraph_format.left_indent = Inches(0.5)
    desc_p.paragraph_format.space_after = Pt(12)
    desc_p.paragraph_format.line_spacing = 1.15
    desc_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_bullet(doc, text, indent=0.5):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

# ==========================================
# CONTENT CONSTRUCTION
# ==========================================

# Add Title
title = doc.add_heading('HearAlert Project: System Architecture & Implementation Steps', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = 'Arial'
    run.bold = True

doc.add_paragraph() # Spacing

# 1. Introduction
add_heading(doc, '1. Project Overview & Scope', level=1)
add_paragraph(doc, "HearAlert is an advanced, real-time audio classification mobile application aimed at providing critical acoustic awareness. By utilizing a highly optimized, edge-based machine learning model (TensorFlow Lite), HearAlert intelligently identifies essential environmental, emergency, and human sounds instantly. These classifications are then translated into priority-based visual cues and customizable haptic feedback routines directly on the user's mobile device.")

# 2. Part 1: Machine Learning Implementation Pipeline (Step-by-Step)
add_heading(doc, '2. Backend: Machine Learning Implementation Pipeline', level=1)
add_paragraph(doc, "The following sequence outlines the precise data engineering and model training pipeline executed to produce the intelligent acoustic backend for the HearAlert application.")

add_step(doc, 1, "Data Aggregation & Corpus Structuring", 
             "The system begins by merging robust foundational datasets (such as the ESC-50) with pre-processed, categorized raw audio data. Concurrently, a synthesis engine dynamically constructs audio files (e.g., oscillating smoke alarms or car alarms) to fill gaps in the original real-world corpus. This results in an extensive training set of over 19,000 files spread entirely across 33 distinct operational categories.")

add_step(doc, 2, "Validation & Test Splitting", 
             "To ensure mathematical validity and to prevent neural network overfitting, the aggregate dataset is partitioned precisely into an 80% Training Set, a 10% Validation Set (used to tune the network autonomously), and a 10% Testing reserve.")

add_step(doc, 3, "Transfer Learning Feature Extraction (YAMNet)", 
             "Instead of analyzing raw spectrograms via a heavy CNN, HearAlert parses 1-second audio frames through Google's pre-trained YAMNet model. YAMNet instantaneously extracts 1024-dimensional feature embeddings, establishing a condensed mathematical blueprint of the audio signal ready for rapid classification.")

add_step(doc, 4, "On-The-Fly Data Augmentation", 
             "During training iterations, incoming audio vectors are dynamically manipulated to simulate severe real-world interference. Operations include 50% probability Gaussian Noise Injection, Time Shifting across the 1-second boundary, Volume Gain Scaling, and minor Pitch Speed modifications. This ensures the model learns to identify patterns underneath substantial environment distortion.")

add_step(doc, 5, "Deep Neural Network (DNN) Custom Top-Layer Training", 
             "The extracted embeddings are fed iteratively into a completely custom DNN head designed specifically for HearAlert. The architecture incorporates three aggressively structured Dense layers (operating at 512, 256, and 128 units respectively). Each block integrates strategic Batch Normalization to accelerate learning and Dropout layers to penalize the memorization of outlier files.")

add_step(doc, 6, "Advanced Strategy Callbacks", 
             "The compiler utilizes a Cosine Decay algorithm with an early optimization Warmup phase, smoothing the navigation through the loss landscape. Furthermore, 'Early Stopping' immediately preserves the strongest checkpoint and prevents over-training once structural improvement plateaus for 15 consecutive epochs.")

add_step(doc, 7, "Edge Quantization (TFLite Int8 Compression)", 
             "The finalized architecture is translated via TensorFlow Lite. A representative dataset drives 'Full-Integer Quantization', scaling absolute 32-bit floats into extremely fast 8-bit integers. This compresses the model graph to roughly 700 KB, guaranteeing millisecond inference speeds specifically optimized for limited Android/iOS hardware.")

doc.add_page_break()

# 3. Part 2: Mobile Application Integration (Step-by-Step)
add_heading(doc, '3. Frontend: Mobile Application Structure & Integration', level=1)
add_paragraph(doc, "The following sequence dictates the structured interaction utilized by the Flutter Frontend to execute and visually communicate the trained acoustic intelligence.")

add_step(doc, 8, "Telemetry & Interface Structuring (/lib/screens/)", 
             "The Flutter engine utilizes primary screens (Onboarding, History, Home) to secure system-level hardware permissions (namely microphone, background execution, and notifications). The layout separates telemetry state management firmly from the actual rendering constraints.")

add_step(doc, 9, "Real-Time Classification Processing (HearAlertClassifierService)", 
             "A dedicated background service intercepts the raw PCM audio stream straight from the native device microphone at specifically 16kHz. This stream buffers up to one-second thresholds where it invokes asynchronous, non-blocking hardware threads holding the 700KB TFLite model, extracting prediction arrays instantly.")

add_step(doc, 10, "State Aggregation (SoundProvider)", 
             "The decoupled SoundProvider class manages the active stream status utilizing modern ChangeNotifier patterns. As prediction thresholds clear confidence hurdles continuously passing from the ML runtime, SoundProvider broadcasts atomic state updates downwards to the completely abstracted UI listeners.")

add_step(doc, 11, "Haptic Telemetry (AlertService) & Notifications", 
             "Relying dynamically upon the training pipeline's generated 'categories_config.json', the AlertService maps the latest acoustic detection instantly. This ensures high-priority emergencies trigger aggressive, relentless vibration integers (e.g., fire alarms emitting [0, 500, 100, 500]), mapping the physical response seamlessly to the identified soundscape regardless of visual availability.")


# Save the document
output_path = '/Users/abusaleem/Hearalert-version-1.1/hearalert-1.0.1/hearalertt/HEARALERT_STEP_BY_STEP_DOCUMENTATION.docx'
doc.save(output_path)
print(f"Document successfully created at {output_path}")
