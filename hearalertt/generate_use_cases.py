from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Initialize Document
doc = Document()

# Set Page Margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.name = 'Arial'
        if level == 0:
            run.bold = True

def add_paragraph(doc, text, bold=False, indent=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    if indent > 0:
        p.paragraph_format.left_indent = Inches(indent)
    
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_use_case(doc, number, title, environment, description, haptic_response):
    # Title
    p = doc.add_paragraph()
    run_num = p.add_run(f"Use Case {number}: ")
    run_num.font.name = 'Arial'
    run_num.font.size = Pt(12)
    run_num.bold = True
    
    run_title = p.add_run(title)
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(12)
    run_title.bold = True
    p.paragraph_format.space_after = Pt(4)
    
    # Environment
    env_p = doc.add_paragraph()
    env_label = env_p.add_run("Environment: ")
    env_label.font.name = 'Arial'
    env_label.font.size = Pt(11)
    env_label.bold = True
    
    env_text = env_p.add_run(environment)
    env_text.font.name = 'Arial'
    env_text.font.size = Pt(11)
    env_text.italic = True
    env_p.paragraph_format.left_indent = Inches(0.25)
    env_p.paragraph_format.space_after = Pt(2)

    # Description
    desc_p = doc.add_paragraph()
    desc_text = desc_p.add_run(description)
    desc_text.font.name = 'Arial'
    desc_text.font.size = Pt(11)
    desc_p.paragraph_format.left_indent = Inches(0.25)
    desc_p.paragraph_format.space_after = Pt(4)
    desc_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Haptic Response
    hap_p = doc.add_paragraph()
    hap_label = hap_p.add_run("System Response: ")
    hap_label.font.name = 'Arial'
    hap_label.font.size = Pt(11)
    hap_label.bold = True
    
    hap_text = hap_p.add_run(haptic_response)
    hap_text.font.name = 'Arial'
    hap_text.font.size = Pt(11)
    hap_p.paragraph_format.left_indent = Inches(0.25)
    hap_p.paragraph_format.space_after = Pt(16)


# ==========================================
# CONTENT CONSTRUCTION
# ==========================================

add_heading(doc, 'HearAlert: Real-Time Operational Use Cases', 0)
doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph() 

add_heading(doc, 'Overview of Real-Time Edge Processing', level=1)
add_paragraph(doc, "HearAlert operates fundamentally as a \"Zero-Latency Edge System.\" Because the machine learning model (TFLite) is fully embedded within the mobile application, the system processes live environmental audio offline. It continuously buffers 1-second audio frames directly from the device microphone at 16kHz, executing rapid inferences without routing data to a cloud server. This strictly preserves user privacy while ensuring life-saving alerts are triggered in milliseconds.")

add_heading(doc, 'Core Operational Scenarios', level=1)
add_paragraph(doc, "Below are structured, real-world implementations demonstrating how HearAlert's 33-class classification engine integrates directly into daily life to protect and inform users who are deaf, hard of hearing, or utilizing noise-canceling headsets.")

# USE CASE 1
add_use_case(doc, 1, "Immediate Emergency Threat Detection", 
             "Urban Outdoors / Indoor Public Spaces",
             "As the user navigates a city street or rests inside an apartment building, unexpected emergency sirens (ambulance, police, fire engines) or localized fire alarms trigger nearby. The user's visual attention is directed elsewhere, and ambient noise levels are high.",
             "The 1-second buffer instantly hits a >80% confidence threshold for 'Siren' or 'Smoke_Alarm'. HearAlert overrides the system UI, driving a critical, relentless haptic vibration pattern (e.g., [1000ms pulse, 500ms rest]) demanding immediate physical attention, accompanied by a bright red visual strobe on the smartwatch or phone screen.")

# USE CASE 2
add_use_case(doc, 2, "Parental Monitoring & Baby Cry Detection", 
             "Residential Home / Nighttime",
             "A deaf parent is sleeping in a separate room from their infant. Standard baby monitors only transmit visual or basic amplified audio signals which are ineffective while the parent is asleep.",
             "The device microphone continuously monitors the ambient bedroom environment. Upon detecting a 'Baby_Cry', HearAlert triggers a distinct, rapid 'heartbeat' haptic vibration pattern designed to wake the parent without causing extreme alarm. A customized push notification is logged to the smartwatch.")

# USE CASE 3
add_use_case(doc, 3, "Navigational Safety & Traffic Awareness", 
             "Pedestrian Commutes / Cyclying",
             "The user is walking near a busy intersection or cycling with fully enclosed noise-canceling headphones. A vehicle approaches rapidly from a blind spot and initiates an aggressive car horn or heavy engine rev.",
             "Instantly identifying 'Car_Horn' or 'Traffic', HearAlert issues a sharp, high-priority burst vibration (e.g., [0, 800, 200, 800]) warning the user of immediate proximal danger, prompting them to visually check their blind spots before stepping into traffic.")

# USE CASE 4
add_use_case(doc, 4, "Household Utility & Social Awareness", 
             "Living Room / Kitchen",
             "The user is watching television or cooking. Visitors arrive and knock on the front door or ring the doorbell. Concurrently, a microwave finishes its cycle out of sight.",
             "The background listener parses the overlapping events. Upon detecting 'Door_Knock' or 'Doorbell', a mid-priority, rhythmic 'Ding-Dong' or 'Knock-Knock' vibration template is sent to the phone. The user knows instantly that someone is at the door without needing a hardwired specialized flashing light system.")

# USE CASE 5
add_use_case(doc, 5, "Background Noise Rejection", 
             "Outdoor Weather / Crowded Environments",
             "The user is walking through heavy rain, high wind, or a crowded environment with intense background chatter. No immediate priority events are occurring, but the microphone is saturated.",
             "The model accurately maps these audio frames to the 'Background' or 'Speech' fallback classes. HearAlert intentionally suppresses haptic feedback, preventing 'notification fatigue.' The system remains silently vigilant, only interrupting the user when an anomaly (like breaking glass or a siren) pierces the noise floor.")


# Save the document
output_path = '/Users/abusaleem/Hearalert-version-1.1/hearalert-1.0.1/hearalertt/HEARALERT_REALTIME_USE_CASES.docx'
doc.save(output_path)
print(f"Document successfully created at {output_path}")
