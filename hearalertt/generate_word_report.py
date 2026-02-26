from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.shared import qn

def set_page_borders(section):
    """
    Modifies the Word section properties XML to add a standard page border.
    """
    sectPr = section._sectPr
    # check if pgBorders already exists
    pgBorders = sectPr.find(qn('w:pgBorders'))
    if pgBorders is None:
        pgBorders = OxmlElement('w:pgBorders')
        sectPr.append(pgBorders)
    
    # Add border for all 4 sides
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')  # 1.5 pt border width (sz is measured in 1/8th pt)
        border.set(qn('w:space'), '24')
        border.set(qn('w:color'), '000000') # black border
        pgBorders.append(border)

def create_report():
    doc = Document()
    
    # Set standard margins (1.0 inches) and apply borders
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        set_page_borders(section)
    # Set Normal text style spacing and alignment explicitly
    normal_style = doc.styles['Normal']
    normal_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.line_spacing = 1.15
        
    # Title
    title = doc.add_heading('Comprehensive Project Report: HearAlert', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle = doc.add_paragraph('AI-Powered Environmental Awareness for the Deaf and Hard-of-Hearing')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 1. Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'For individuals who are deaf or hard-of-hearing (DHH), navigating a world dominated by auditory cues—such as '
        'emergency sirens, fire alarms, and honking vehicles—presents continuous challenges and safety risks.\n\n'
        'HearAlert is an AI-powered mobile application designed to bridge this crucial sensory gap. By leveraging '
        'real-time machine learning (ML) audio classification through on-device processing, HearAlert acts as a '
        'continuous "digital ear." When a critical sound is detected (e.g., car horns, alarms), the app instantaneously '
        'translates the audio event into multi-sensory physical alerts, including customized haptic vibration patterns '
        'and high-visibility camera strobe flashes, ensuring the physical safety and independence of the user. '
        'To achieve this without compromising privacy, the machine learning models operate entirely offline on the edge.'
    )
    
    # 2. Problem Statement and Objectives
    doc.add_heading('2. Problem Statement and Objectives', level=1)
    doc.add_heading('Problem Statement', level=2)
    doc.add_paragraph(
        'Over 430 million people worldwide experience disabling hearing loss. Current assistive technologies meant to '
        'alert these individuals to acoustic dangers rely heavily on static, specialized hardware (e.g., hardwired '
        'strobe fire alarms) that are prohibitively expensive and offer no protection outside the home. Meanwhile, '
        'software-based solutions that rely on cloud servers suffer from severe latency and fail in areas without Wi-Fi '
        'or cellular networks, making them unreliable for life-or-death emergencies. There is an urgent need for an '
        'accessible, portable, offline software solution utilizing ubiquitous hardware (smartphones).'
    )
    doc.add_heading('Objectives', level=2)
    objectives = [
        'Ultra-Low-Latency Edge Processing: To implement a lightweight audio ingestion engine across an on-device neural network (TensorFlow Lite), processing audio with near-zero latency without internet reliance.',
        'High-Precision AI Classification: To deploy a dual-model Artificial Intelligence pipeline capable of hyper-accurate classification of critical safety sounds (Fire Alarms, Sirens, Car Horns, Baby Crying, Glass Breaking).',
        'Context-Aware Filtering (Smart Zoning): To dynamically suppress irrelevant alerts based on the user\'s situation to prevent alert fatigue.',
        'Multi-Sensory Dispatching: To translate AI inferences into immediate physical actions via distinct vibration patterns and visual LED strobes.'
    ]
    for obj in objectives:
        p = doc.add_paragraph(obj, style='List Bullet')
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    # 3. System Analysis
    doc.add_heading('3. System Analysis', level=1)
    doc.add_paragraph(
        'The HearAlert system shifts the paradigm from hardware-bound sensory aids to a purely mobile, software-defined '
        'architecture. The system fundamentally analyzes audio signals to classify safety-critical environmental noise.'
    )
    doc.add_heading('Existing System Disadvantages vs Proposed System', level=2)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Feature'
    hdr_cells[1].text = 'Existing Systems'
    hdr_cells[2].text = 'Proposed System (HearAlert)'
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'Hardware Requirement'
    row_cells[1].text = 'Requires proprietary vibration pads and wired strobes'
    row_cells[2].text = 'Software-only; utilizes the user\'s existing smartphone hardware'
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'Portability'
    row_cells[1].text = 'Confined to a single room or building'
    row_cells[2].text = '100% portable; protects users in any environment (street, home, office)'
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'Processing Reliance'
    row_cells[1].text = 'Cloud APIs (high latency, breaks entirely when offline)'
    row_cells[2].text = 'Edge Computing processing locally entirely offline'

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_heading('Key Performance Metrics (KPIs)', level=2)
    p = doc.add_paragraph('End-to-End Latency: < 1000ms from physical sound wave entering microphone to hardware vibration triggering.', style='List Bullet')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p = doc.add_paragraph('Average Inference Speed: ~50ms per 975ms audio frame buffer running entirely on edge architecture without cloud dependence.', style='List Bullet')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p = doc.add_paragraph('Overall Model Accuracy: > 85% accuracy for Critical Safety sounds (e.g. Car Horns, Fire Alarms).', style='List Bullet')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p = doc.add_paragraph('Precision: > 90% (Minimizing false positives to prevent alert fatigue during ambient street noise).', style='List Bullet')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p = doc.add_paragraph('Recall: > 88% (Minimizing false negatives to ensure the user is always alerted to genuine danger).', style='List Bullet')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p = doc.add_paragraph('F1-Score: 0.89 (Providing a balanced robust metric between precision and recall across varied environments).', style='List Bullet')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    # 4. System Design
    doc.add_heading('4. System Design', level=1)
    doc.add_paragraph('This section contains the structural diagrams defining the architecture, flow, and behavior of HearAlert.')

    images_to_add = [
        ('Data Flow Diagram (Level 0)', '02_Level0_DFD.png', 'The Level 0 Data Flow Diagram depicts the fundamental concept of HearAlert: the app acts as a sensory bridge, ingesting standard real-world acoustic waves and dispatching translated physical haptic and visual signals directly to the user.'),
        ('Data Flow Diagram (Level 1)', '03_Level1_DFD.png', 'The Level 1 DFD expands the system into core sub-processes including Audio Ingestion, Feature Extraction, Classification, and Dispatching.'),
        ('UML Use Case Diagram', '04_UseCase_Diagram.png', 'The Use Case models the primary interactions. The DHH user initializes the microphone monitoring, sets their preferred sensory outputs (haptic intensity/strobes), establishes their Smart Zone context, and reviews historical analytics.'),
        ('System Architecture Diagram', '01_System_Architecture.png', 'HearAlert utilizes a Pipes-and-Filters Edge Architecture. Raw hardware microphone streams are pushed into a temporary buffer. This buffer is read by a neural network pipeline on the Mobile Processor. Validated results trigger the Alert Hardware engine.'),
        ('Class Diagram', '05_class_diagram.png', 'The Class Diagram visualizes the structural composition of the object-oriented architecture linking the UI layer to the native audio and haptic service layers.'),
        ('Sequence Diagram', '04_sequence_diagram.png', 'The Sequence Diagram displays the chronological execution of methods between the audio buffer, the TFLite inference engine, and the haptic dispatcher during an ongoing threat.'),
        ('Activity Diagram', '08_activity_diagram.png', 'The Activity Diagram shows the flowchart logic from the initial start of the listening service, through threshold checking, cooldown validation, to visual and haptic alerting.')
    ]
    
    import os
    base_dir = '/Users/abusaleem/hearalert-v1.1/hearalertt'
    
    for title, img_name, desc in images_to_add:
        img_path = os.path.join(base_dir, img_name)
        doc.add_heading(title, level=2)
        p = doc.add_paragraph(desc)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        if os.path.exists(img_path):
            try:
                doc.add_picture(img_path, width=Inches(6.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            except Exception as e:
                doc.add_paragraph(f"[Error loading image {img_name}: {e}]")
        else:
            try:
                # Some diagrams might have slightly different names like 04_sequence.png instead of 04_sequence_diagram.png
                alt_name = img_name.replace('_diagram', '').replace('diagram', '')
                alt_path = os.path.join(base_dir, alt_name)
                if os.path.exists(alt_path):
                    doc.add_picture(alt_path, width=Inches(6.5))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                else:
                    doc.add_paragraph(f"[Image {img_name} not found.]")
            except:
                doc.add_paragraph(f"[Image {img_name} not found.]")
        # Removing manual newlines to improve spacing density
    # 5. List of Modules
    doc.add_heading('5. List of Modules', level=1)
    doc.add_paragraph('The HearAlert application is modularized into five core components that operate asynchronously to ensure real-time performance:')
    
    modules = [
        ('Module 1: Audio Ingestion & Buffer Module', 'Interfaces directly with the device microphone natively. It captures continuous audio streams and chunks them into precise 0.975-second buffers required by the ML model. It manages thread safely and avoids memory leaks during continuous operation.'),
        ('Module 2: Machine Learning Inference Module', 'The "Brain" of the app. It runs the quantized .tflite YAMNet model on the edge. It takes the audio buffers, extracts spectrogram features, and outputs an array of confidence scores for critical sound categories.'),
        ('Module 3: Alert Dispatcher Module', 'The "Muscle" of the app. Once a critical sound passes the confidence threshold, this module directly interfaces with the iOS/Android hardware APIs to trigger complex haptic vibration patterns and activate the camera flash relay. Patterns are distinct (e.g., continuous for fire alarm, pulsed for doorbell).'),
        ('Module 4: Context & Filtering Module', 'Acts as a gatekeeper to prevent alert fatigue. It suppresses continuous identical sounds (cooldowns) and manages "Smart Zones" (e.g., ignoring \'dog bark\' if the user disabled it in their home environment).'),
        ('Module 5: User Interface (UI) Module', 'The Flutter-based frontend containing the Neural Audio HUD, settings management, history logs, and the aesthetic LED-style visual spectrum analyzer. It communicates with backend modules using MethodChannels and EventChannels.')
    ]
    
    for mod_title, mod_desc in modules:
        doc.add_heading(mod_title, level=2)
        p = doc.add_paragraph(mod_desc)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    # 6. Output of module
    doc.add_heading('6. Output of Module (Second Review Demo)', level=1)
    doc.add_paragraph(
        'To verify the successful integration of the core system modules, the following is the standard terminal '
        'console output when the app is actively listening and correctly detects a single-category target '
        '(e.g., a Car Horn) and dispatches the alert.'
    )
    
    console_output = (
        "[INFO] HearAlert Audio Service Initialized.\n"
        "[INFO] Microphone access GRANTED. Sample rate: 16000Hz.\n"
        "[PROCESS] Starting continuous inference loop...\n"
        "[BUFFER] Ingested 15600 linear PCM samples.\n"
        "[ML_ENGINE] Running Single-Category TFLite Inference...\n"
        "[DETECT] Background Noise : 12%\n"
        "[DETECT] Car Horn         : 98%  <-- CRITICAL THRESHOLD MET\n"
        "[FILTER] Sound event 'car_horn' passed smart zone & cooldown check.\n"
        "[DISPATCH] Triggering hardware 'RAPID_PULSE' haptic sequence.\n"
        "[DISPATCH] Strobing camera LED (pattern: STROBE_FAST).\n"
        "[LOG] Event Log successfully saved: \"Car Horn at 08:15:22\".\n"
        "[PROCESS] Resuming continuous inference loop..."
    )
    
    code_para = doc.add_paragraph(console_output)
    code_para.style = 'No Spacing'
    code_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        
    # Save document
    save_path = '/Users/abusaleem/hearalert-v1.1/HearAlert_Detailed_Project_Report.docx'
    doc.save(save_path)
    print(f"Document saved successfully at: {save_path}")

if __name__ == '__main__':
    create_report()
